from  docx import Document
from pathlib import Path

type_panes = { "quantity": None, "type_glass": None, "width": None, "height": None, "oa": None }



    
def parse_docx(file_path):
    # Открываем документ .docx
    doc = Document(file_path)
    
    # Переменная для хранения всего текста
    full_text = []
    
    # Извлечение текста из параграфов
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    # Соединяем все параграфы в один текст с разделением по строкам
    return '\n'.join(full_text)

def parse_docx_with_tables(file_path):
    # Открытие документа
    doc = Document(file_path)
    
    # Список для хранения текста
    full_text = []
    table_row = {"Position": None, "QNTITY":None, "width": None, "height":None, "OA":None, "Muntin Bars":None}
    
    # Извлечение текста из параграфов
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    print(full_text)

    # Извлечение текста из таблиц
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                if cell.text != ' ':
                    row_data.append(cell.text)
            full_text.append('\t'.join(row_data))  # Добавляем строку таблицы, разделенную табуляцией

    return '\n'.join(full_text)

def get_data_table(test):
    keyword = "MUNTIN BARS"
    index = test.find(keyword)
    if index !=-1:
        result = test[index + len(keyword):].strip()
        return(result)  # Output: "language for beginners and professionals."
    else:
        print("Keyword not found.")


def fill_dict(type_panes, table):
    for line in table.splitlines():
        parts = line.split('\t')
        return parts
        '''
        type_panes["quantity"] = parts[0].strip()
        type_panes["type_glass"] = parts[1].strip()
        type_panes["width"] = parts[3].strip()
        type_panes["height"] = parts[4].strip()
        type_panes["oa"] = parts[5].strip()
        '''
    
    return type_panes
# Использование функции для парсинга
file_path = Path('/home/roma/parsing') / 'AlexLinkmeyer.docx' # Замените на путь к вашему файлу
text = parse_docx_with_tables(file_path)
table_in_text = get_data_table(text) 

# Вывод текста в консоль
pure_dict = fill_dict(type_panes, table_in_text)




